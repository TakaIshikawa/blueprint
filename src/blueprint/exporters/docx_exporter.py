<<<<<<< HEAD
"""Word DOCX exporter with editable tables and charts."""
=======
"""Word DOCX exporter for execution plans."""
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM

from __future__ import annotations

import io
<<<<<<< HEAD
import zipfile
from datetime import datetime, timezone
from typing import Any
from xml.sax.saxutils import escape as xml_escape
=======
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM

from blueprint.exporters.base import TargetExporter

# ---------------------------------------------------------------------------
<<<<<<< HEAD
# OOXML namespace constants
# ---------------------------------------------------------------------------

_NS_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"
_NS_PKG_RELS = "http://schemas.openxmlformats.org/package/2006/relationships"
_NS_DOC_RELS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_CP = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
_NS_DC = "http://purl.org/dc/elements/1.1/"
_NS_DCTERMS = "http://purl.org/dc/terms/"
_NS_XSI = "http://www.w3.org/2001/XMLSchema-instance"

# ---------------------------------------------------------------------------
# Style definitions
=======
# Constants
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
# ---------------------------------------------------------------------------

STATUS_LABELS: dict[str, str] = {
    "pending": "Pending",
    "in_progress": "In Progress",
    "completed": "Completed",
    "blocked": "Blocked",
    "skipped": "Skipped",
}

<<<<<<< HEAD
STATUS_COLORS: dict[str, str] = {
    "pending": "CC9900",
    "in_progress": "0066CC",
    "completed": "008000",
    "blocked": "CC0000",
    "skipped": "777777",
}

STYLE_CONFIGS: dict[str, dict[str, str]] = {
    "default": {
        "primary": "003366",
        "secondary": "336699",
        "accent": "CC6600",
        "header_bg": "003366",
        "header_text": "FFFFFF",
        "stripe_bg": "EBF1F7",
        "border": "CCCCCC",
    },
    "corporate": {
        "primary": "1B365D",
        "secondary": "4A7298",
        "accent": "8B6914",
        "header_bg": "1B365D",
        "header_text": "FFFFFF",
        "stripe_bg": "E8EDF2",
        "border": "B0B0B0",
    },
    "modern": {
        "primary": "2D3748",
        "secondary": "4A5568",
        "accent": "3182CE",
        "header_bg": "2D3748",
        "header_text": "FFFFFF",
        "stripe_bg": "EDF2F7",
        "border": "CBD5E0",
    },
}


def _count_statuses(tasks: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for task in tasks:
        status = task.get("status") or "pending"
        counts[status] = counts.get(status, 0) + 1
    return counts


# ---------------------------------------------------------------------------
# OOXML XML generators
# ---------------------------------------------------------------------------


def _content_types_xml() -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="{_NS_CT}">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
  <Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>
  <Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>
</Types>"""


def _rels_xml() -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{_NS_PKG_RELS}">
  <Relationship Id="rId1" Type="{_NS_DOC_RELS}/officeDocument" Target="word/document.xml"/>
  <Relationship Id="rId2" Type="{_NS_PKG_RELS}/metadata/core-properties" Target="docProps/core.xml"/>
</Relationships>"""


def _word_rels_xml() -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="{_NS_PKG_RELS}">
  <Relationship Id="rId1" Type="{_NS_DOC_RELS}/styles" Target="styles.xml"/>
  <Relationship Id="rId2" Type="{_NS_DOC_RELS}/numbering" Target="numbering.xml"/>
</Relationships>"""


def _core_props_xml(
    title: str = "",
    author: str = "",
    subject: str = "",
    keywords: str = "",
) -> str:
    now = datetime.now(tz=timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<cp:coreProperties xmlns:cp="{_NS_CP}" xmlns:dc="{_NS_DC}" xmlns:dcterms="{_NS_DCTERMS}" xmlns:xsi="{_NS_XSI}">
  <dc:title>{xml_escape(title)}</dc:title>
  <dc:creator>{xml_escape(author)}</dc:creator>
  <dc:subject>{xml_escape(subject)}</dc:subject>
  <cp:keywords>{xml_escape(keywords)}</cp:keywords>
  <dcterms:created xsi:type="dcterms:W3CDTF">{now}</dcterms:created>
  <dcterms:modified xsi:type="dcterms:W3CDTF">{now}</dcterms:modified>
</cp:coreProperties>"""


def _styles_xml(style_config: dict[str, str]) -> str:
    primary = style_config.get("primary", "003366")
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="{_NS_W}">
  <w:docDefaults>
    <w:rPrDefault>
      <w:rPr>
        <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri" w:cs="Calibri"/>
        <w:sz w:val="22"/>
        <w:szCs w:val="22"/>
      </w:rPr>
    </w:rPrDefault>
    <w:pPrDefault>
      <w:pPr>
        <w:spacing w:after="120" w:line="276" w:lineRule="auto"/>
      </w:pPr>
    </w:pPrDefault>
  </w:docDefaults>
  <w:style w:type="paragraph" w:styleId="Normal" w:default="1">
    <w:name w:val="Normal"/>
    <w:rPr><w:sz w:val="22"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:pPr>
      <w:keepNext/>
      <w:keepLines/>
      <w:spacing w:before="360" w:after="120"/>
      <w:outlineLvl w:val="0"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
      <w:b/>
      <w:color w:val="{primary}"/>
      <w:sz w:val="36"/>
      <w:szCs w:val="36"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:pPr>
      <w:keepNext/>
      <w:keepLines/>
      <w:spacing w:before="240" w:after="80"/>
      <w:outlineLvl w:val="1"/>
    </w:pPr>
    <w:rPr>
      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>
      <w:b/>
      <w:color w:val="{primary}"/>
      <w:sz w:val="28"/>
      <w:szCs w:val="28"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading3">
    <w:name w:val="heading 3"/>
    <w:basedOn w:val="Normal"/>
    <w:next w:val="Normal"/>
    <w:pPr>
      <w:keepNext/>
      <w:keepLines/>
      <w:spacing w:before="200" w:after="60"/>
      <w:outlineLvl w:val="2"/>
    </w:pPr>
    <w:rPr>
      <w:b/>
      <w:color w:val="{primary}"/>
      <w:sz w:val="24"/>
      <w:szCs w:val="24"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Title">
    <w:name w:val="Title"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:spacing w:after="200"/>
      <w:jc w:val="center"/>
    </w:pPr>
    <w:rPr>
      <w:b/>
      <w:color w:val="{primary}"/>
      <w:sz w:val="52"/>
      <w:szCs w:val="52"/>
    </w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Subtitle">
    <w:name w:val="Subtitle"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:jc w:val="center"/></w:pPr>
    <w:rPr>
      <w:color w:val="666666"/>
      <w:sz w:val="28"/>
      <w:szCs w:val="28"/>
    </w:rPr>
  </w:style>
  <w:style w:type="table" w:styleId="TableGrid">
    <w:name w:val="Table Grid"/>
    <w:tblPr>
      <w:tblBorders>
        <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
      </w:tblBorders>
    </w:tblPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="ListBullet">
    <w:name w:val="List Bullet"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr>
      <w:numPr>
        <w:numId w:val="1"/>
      </w:numPr>
      <w:ind w:left="720"/>
    </w:pPr>
  </w:style>
</w:styles>"""


def _numbering_xml() -> str:
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="{_NS_W}">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="\u2022"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol" w:hint="default"/></w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>"""


# ---------------------------------------------------------------------------
# Document body builders
# ---------------------------------------------------------------------------


def _para(
    text: str,
    style: str = "Normal",
    bold: bool = False,
    color: str = "",
    size: int = 0,
    alignment: str = "",
) -> str:
    """Build a paragraph XML element."""
    ppr_parts: list[str] = []
    if style != "Normal":
        ppr_parts.append(f'<w:pStyle w:val="{style}"/>')
    if alignment:
        ppr_parts.append(f'<w:jc w:val="{alignment}"/>')
    ppr = f"<w:pPr>{' '.join(ppr_parts)}</w:pPr>" if ppr_parts else ""

    rpr_parts: list[str] = []
    if bold:
        rpr_parts.append("<w:b/>")
    if color:
        rpr_parts.append(f'<w:color w:val="{color}"/>')
    if size:
        rpr_parts.append(f'<w:sz w:val="{size}"/><w:szCs w:val="{size}"/>')
    rpr = f"<w:rPr>{' '.join(rpr_parts)}</w:rPr>" if rpr_parts else ""

    return f"<w:p>{ppr}<w:r>{rpr}<w:t xml:space=\"preserve\">{xml_escape(text)}</w:t></w:r></w:p>"


def _page_break() -> str:
    return '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'


def _table_cell(
    text: str,
    bg_color: str = "",
    text_color: str = "",
    bold: bool = False,
    width: int = 0,
) -> str:
    """Build a table cell XML element."""
    tc_pr_parts: list[str] = []
    if bg_color:
        tc_pr_parts.append(f'<w:shd w:val="clear" w:color="auto" w:fill="{bg_color}"/>')
    if width:
        tc_pr_parts.append(f'<w:tcW w:w="{width}" w:type="dxa"/>')
    tc_pr = f"<w:tcPr>{' '.join(tc_pr_parts)}</w:tcPr>" if tc_pr_parts else ""

    rpr_parts: list[str] = []
    if bold:
        rpr_parts.append("<w:b/>")
    if text_color:
        rpr_parts.append(f'<w:color w:val="{text_color}"/>')
    rpr = f"<w:rPr>{' '.join(rpr_parts)}</w:rPr>" if rpr_parts else ""

    return (
        f"<w:tc>{tc_pr}"
        f'<w:p><w:r>{rpr}<w:t xml:space="preserve">{xml_escape(text)}</w:t></w:r></w:p>'
        f"</w:tc>"
    )


def _table_row(cells: list[str]) -> str:
    return f"<w:tr>{''.join(cells)}</w:tr>"


def _build_task_table(
    tasks: list[dict[str, Any]], style_config: dict[str, str]
) -> str:
    """Build the main task breakdown table."""
    header_bg = style_config.get("header_bg", "003366")
    header_text = style_config.get("header_text", "FFFFFF")
    stripe_bg = style_config.get("stripe_bg", "EBF1F7")

    headers = ["ID", "Title", "Status", "Milestone", "Owner", "Complexity", "Est. Hours"]
    col_widths = [1200, 2400, 1200, 1400, 1200, 1100, 1000]

    header_cells = [
        _table_cell(h, bg_color=header_bg, text_color=header_text, bold=True, width=w)
        for h, w in zip(headers, col_widths)
    ]

    rows = [_table_row(header_cells)]

    for i, task in enumerate(tasks):
        task_id = task.get("id") or ""
        title = task.get("title") or ""
        status = task.get("status") or "pending"
        status_label = STATUS_LABELS.get(status, status)
        milestone = task.get("milestone") or "-"
        owner = task.get("owner_type") or "-"
        complexity = task.get("estimated_complexity") or "-"
        hours = task.get("estimated_hours")
        hours_str = f"{hours:.1f}" if hours else "-"

        bg = stripe_bg if i % 2 == 0 else ""
        status_color = STATUS_COLORS.get(status, "000000")

        cells = [
            _table_cell(task_id, bg_color=bg, width=col_widths[0]),
            _table_cell(title, bg_color=bg, width=col_widths[1]),
            _table_cell(status_label, bg_color=bg, text_color=status_color, bold=True, width=col_widths[2]),
            _table_cell(milestone, bg_color=bg, width=col_widths[3]),
            _table_cell(owner, bg_color=bg, width=col_widths[4]),
            _table_cell(complexity, bg_color=bg, width=col_widths[5]),
            _table_cell(hours_str, bg_color=bg, width=col_widths[6]),
        ]
        rows.append(_table_row(cells))

    border = style_config.get("border", "CCCCCC")
    return f"""<w:tbl>
<w:tblPr>
  <w:tblStyle w:val="TableGrid"/>
  <w:tblW w:w="9500" w:type="dxa"/>
  <w:tblBorders>
    <w:top w:val="single" w:sz="4" w:space="0" w:color="{border}"/>
    <w:left w:val="single" w:sz="4" w:space="0" w:color="{border}"/>
    <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border}"/>
    <w:right w:val="single" w:sz="4" w:space="0" w:color="{border}"/>
    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border}"/>
    <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{border}"/>
  </w:tblBorders>
</w:tblPr>
<w:tblGrid>
  {''.join(f'<w:gridCol w:w="{w}"/>' for w in col_widths)}
</w:tblGrid>
{''.join(rows)}
</w:tbl>"""


def _build_dependency_table(
    tasks: list[dict[str, Any]], style_config: dict[str, str]
) -> str:
    """Build the dependency matrix table."""
    header_bg = style_config.get("header_bg", "003366")
    header_text = style_config.get("header_text", "FFFFFF")
    stripe_bg = style_config.get("stripe_bg", "EBF1F7")

    task_map = {t.get("id") or "": t for t in tasks}
    deps_found = False

    header_cells = [
        _table_cell("Source Task", bg_color=header_bg, text_color=header_text, bold=True, width=3000),
        _table_cell("Depends On", bg_color=header_bg, text_color=header_text, bold=True, width=3000),
        _table_cell("Status", bg_color=header_bg, text_color=header_text, bold=True, width=1500),
    ]
    rows = [_table_row(header_cells)]

    i = 0
    for task in tasks:
        deps = task.get("depends_on") or []
        if not deps:
            continue
        deps_found = True
        task_id = task.get("id") or ""
        task_title = task.get("title") or ""

        for dep_id in deps:
            dep_task = task_map.get(dep_id, {})
            dep_title = dep_task.get("title") or dep_id
            dep_status = dep_task.get("status") or "unknown"
            dep_status_label = STATUS_LABELS.get(dep_status, dep_status)
            status_color = STATUS_COLORS.get(dep_status, "000000")

            bg = stripe_bg if i % 2 == 0 else ""
            cells = [
                _table_cell(f"{task_id}: {task_title}", bg_color=bg, width=3000),
                _table_cell(f"{dep_id}: {dep_title}", bg_color=bg, width=3000),
                _table_cell(dep_status_label, bg_color=bg, text_color=status_color, bold=True, width=1500),
            ]
            rows.append(_table_row(cells))
            i += 1

    if not deps_found:
        return _para("No dependencies defined.", color="777777")

    return f"""<w:tbl>
<w:tblPr>
  <w:tblStyle w:val="TableGrid"/>
  <w:tblW w:w="7500" w:type="dxa"/>
</w:tblPr>
<w:tblGrid>
  <w:gridCol w:w="3000"/>
  <w:gridCol w:w="3000"/>
  <w:gridCol w:w="1500"/>
</w:tblGrid>
{''.join(rows)}
</w:tbl>"""


def _build_risk_table(
    tasks: list[dict[str, Any]],
    brief: dict[str, Any],
    style_config: dict[str, str],
) -> str:
    """Build risk register table."""
    header_bg = style_config.get("header_bg", "003366")
    header_text = style_config.get("header_text", "FFFFFF")
    stripe_bg = style_config.get("stripe_bg", "EBF1F7")

    risk_color_map = {"high": "CC0000", "medium": "CC9900", "low": "008000"}

    header_cells = [
        _table_cell("Risk", bg_color=header_bg, text_color=header_text, bold=True, width=4000),
        _table_cell("Level", bg_color=header_bg, text_color=header_text, bold=True, width=1500),
        _table_cell("Source", bg_color=header_bg, text_color=header_text, bold=True, width=2000),
    ]
    rows = [_table_row(header_cells)]

    i = 0
    # Risks from brief
    for risk in brief.get("risks") or []:
        bg = stripe_bg if i % 2 == 0 else ""
        cells = [
            _table_cell(risk, bg_color=bg, width=4000),
            _table_cell("Identified", bg_color=bg, text_color="CC6600", bold=True, width=1500),
            _table_cell("Brief", bg_color=bg, width=2000),
        ]
        rows.append(_table_row(cells))
        i += 1

    # High-risk tasks
    for task in tasks:
        risk_level = task.get("risk_level") or ""
        if risk_level in ("high", "medium"):
            bg = stripe_bg if i % 2 == 0 else ""
            color = risk_color_map.get(risk_level, "000000")
            cells = [
                _table_cell(f"Task: {task.get('title', '')}", bg_color=bg, width=4000),
                _table_cell(risk_level.capitalize(), bg_color=bg, text_color=color, bold=True, width=1500),
                _table_cell(f"Task {task.get('id', '')}", bg_color=bg, width=2000),
            ]
            rows.append(_table_row(cells))
            i += 1

    if i == 0:
        return _para("No risks identified.", color="777777")

    return f"""<w:tbl>
<w:tblPr>
  <w:tblStyle w:val="TableGrid"/>
  <w:tblW w:w="7500" w:type="dxa"/>
</w:tblPr>
<w:tblGrid>
  <w:gridCol w:w="4000"/>
  <w:gridCol w:w="1500"/>
  <w:gridCol w:w="2000"/>
</w:tblGrid>
{''.join(rows)}
</w:tbl>"""


def _build_resource_table(
    tasks: list[dict[str, Any]], style_config: dict[str, str]
) -> str:
    """Build resource allocation table."""
    header_bg = style_config.get("header_bg", "003366")
    header_text = style_config.get("header_text", "FFFFFF")
    stripe_bg = style_config.get("stripe_bg", "EBF1F7")

    owner_data: dict[str, dict[str, Any]] = {}
    for task in tasks:
        owner = task.get("owner_type") or "Unassigned"
        if owner not in owner_data:
            owner_data[owner] = {"tasks": 0, "hours": 0.0, "completed": 0}
        owner_data[owner]["tasks"] += 1
        owner_data[owner]["hours"] += task.get("estimated_hours") or 0
        if task.get("status") == "completed":
            owner_data[owner]["completed"] += 1

    header_cells = [
        _table_cell("Owner", bg_color=header_bg, text_color=header_text, bold=True, width=2000),
        _table_cell("Tasks", bg_color=header_bg, text_color=header_text, bold=True, width=1200),
        _table_cell("Completed", bg_color=header_bg, text_color=header_text, bold=True, width=1500),
        _table_cell("Est. Hours", bg_color=header_bg, text_color=header_text, bold=True, width=1500),
    ]
    rows = [_table_row(header_cells)]

    for i, (owner, data) in enumerate(sorted(owner_data.items())):
        bg = stripe_bg if i % 2 == 0 else ""
        cells = [
            _table_cell(owner, bg_color=bg, width=2000),
            _table_cell(str(data["tasks"]), bg_color=bg, width=1200),
            _table_cell(str(data["completed"]), bg_color=bg, width=1500),
            _table_cell(f"{data['hours']:.1f}", bg_color=bg, width=1500),
        ]
        rows.append(_table_row(cells))

    return f"""<w:tbl>
<w:tblPr>
  <w:tblStyle w:val="TableGrid"/>
  <w:tblW w:w="6200" w:type="dxa"/>
</w:tblPr>
<w:tblGrid>
  <w:gridCol w:w="2000"/>
  <w:gridCol w:w="1200"/>
  <w:gridCol w:w="1500"/>
  <w:gridCol w:w="1500"/>
</w:tblGrid>
{''.join(rows)}
</w:tbl>"""


# ---------------------------------------------------------------------------
# Document body assembly
# ---------------------------------------------------------------------------


def _build_document_body(
    plan: dict[str, Any],
    brief: dict[str, Any],
    style_config: dict[str, str],
) -> str:
    """Assemble the full document.xml body content."""
    tasks = plan.get("tasks") or []
    milestones = plan.get("milestones") or []
    status_counts = _count_statuses(tasks)
    total = sum(status_counts.values())
    completed = status_counts.get("completed", 0)
    progress = round(completed / total * 100) if total > 0 else 0
    total_hours = sum(t.get("estimated_hours") or 0 for t in tasks)
    plan_id = plan.get("id") or "Execution Plan"
    brief_title = brief.get("title") or ""
    status = plan.get("status") or "draft"

    parts: list[str] = []

    # --- Cover page ---
    parts.append(_para("Execution Plan", style="Title"))
    parts.append(_para(plan_id, style="Subtitle"))
    if brief_title:
        parts.append(_para(brief_title, style="Subtitle"))
    engine = plan.get("target_engine") or ""
    repo = plan.get("target_repo") or ""
    if engine:
        parts.append(_para(f"Engine: {engine}", alignment="center", color="666666"))
    if repo:
        parts.append(_para(f"Repository: {repo}", alignment="center", color="666666"))
    date_str = datetime.now(tz=timezone.utc).strftime("%B %d, %Y")
    parts.append(_para(f"Generated: {date_str}", alignment="center", color="999999"))
    parts.append(_para(f"Status: {status.upper()}", alignment="center", bold=True, color=style_config.get("primary", "003366")))
    parts.append(_page_break())

    # --- Executive Summary ---
    parts.append(_para("Executive Summary", style="Heading1"))
    parts.append(_para(f"Plan Status: {status.replace('_', ' ').title()}", bold=True))
    parts.append(_para(f"Progress: {progress}% ({completed}/{total} tasks completed)"))
    parts.append(_para(f"Total Estimated Hours: {total_hours:.0f}"))

    if brief.get("problem_statement"):
        parts.append(_para("Problem Statement", style="Heading2"))
        parts.append(_para(brief["problem_statement"]))

    if brief.get("mvp_goal"):
        parts.append(_para("MVP Goal", style="Heading2"))
        parts.append(_para(brief["mvp_goal"]))

    # Status breakdown
    parts.append(_para("Status Breakdown", style="Heading3"))
    for status_key in ["completed", "in_progress", "pending", "blocked", "skipped"]:
        count = status_counts.get(status_key, 0)
        if count > 0:
            label = STATUS_LABELS.get(status_key, status_key)
            color = STATUS_COLORS.get(status_key, "000000")
            parts.append(_para(f"{label}: {count}", color=color, bold=True))

    parts.append(_page_break())

    # --- Task Breakdown ---
    parts.append(_para("Task Breakdown", style="Heading1"))
    parts.append(_build_task_table(tasks, style_config))
    parts.append(_page_break())

    # --- Timeline ---
    if milestones or tasks:
        parts.append(_para("Timeline", style="Heading1"))
        for ms in milestones:
            ms_name = ms.get("name") or "Milestone"
            ms_desc = ms.get("description") or ""
            parts.append(_para(ms_name, style="Heading2"))
            if ms_desc:
                parts.append(_para(ms_desc))
            # Tasks in this milestone
            ms_tasks = [t for t in tasks if t.get("milestone") == ms_name]
            for task in ms_tasks:
                status_key = task.get("status") or "pending"
                label = STATUS_LABELS.get(status_key, status_key)
                color = STATUS_COLORS.get(status_key, "000000")
                parts.append(
                    _para(
                        f"  [{label}] {task.get('id', '')}: {task.get('title', '')}",
                        color=color,
                    )
                )
        # Tasks without milestone
        no_ms = [t for t in tasks if not t.get("milestone")]
        if no_ms:
            parts.append(_para("Other Tasks", style="Heading2"))
            for task in no_ms:
                status_key = task.get("status") or "pending"
                label = STATUS_LABELS.get(status_key, status_key)
                color = STATUS_COLORS.get(status_key, "000000")
                parts.append(
                    _para(
                        f"  [{label}] {task.get('id', '')}: {task.get('title', '')}",
                        color=color,
                    )
                )
        parts.append(_page_break())

    # --- Dependency Matrix ---
    parts.append(_para("Dependency Matrix", style="Heading1"))
    parts.append(_build_dependency_table(tasks, style_config))
    parts.append(_page_break())

    # --- Risk Register ---
    parts.append(_para("Risk Register", style="Heading1"))
    parts.append(_build_risk_table(tasks, brief, style_config))
    parts.append(_page_break())

    # --- Resource Allocation ---
    parts.append(_para("Resource Allocation", style="Heading1"))
    parts.append(_build_resource_table(tasks, style_config))
    parts.append(_page_break())

    # --- Appendix: Detailed Descriptions ---
    parts.append(_para("Appendix: Detailed Task Descriptions", style="Heading1"))
    for task in tasks:
        task_id = task.get("id") or ""
        title = task.get("title") or ""
        description = task.get("description") or ""
        status_key = task.get("status") or "pending"
        criteria = task.get("acceptance_criteria") or []
        deps = task.get("depends_on") or []

        parts.append(_para(f"{task_id}: {title}", style="Heading2"))
        parts.append(_para(f"Status: {STATUS_LABELS.get(status_key, status_key)}", bold=True, color=STATUS_COLORS.get(status_key, "000000")))
        parts.append(_para(description))

        if deps:
            parts.append(_para(f"Dependencies: {', '.join(str(d) for d in deps)}", color="666666"))

        if criteria:
            parts.append(_para("Acceptance Criteria:", bold=True))
            for ac in criteria:
                parts.append(_para(f"\u2022 {ac}", style="ListBullet"))

    return "\n".join(parts)


def _document_xml(body_content: str) -> str:
    """Wrap body content in the document.xml structure."""
    return f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="{_NS_W}" xmlns:r="{_NS_R}">
<w:body>
{body_content}
<w:sectPr>
  <w:pgSz w:w="12240" w:h="15840"/>
  <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"
           w:header="720" w:footer="720" w:gutter="0"/>
</w:sectPr>
</w:body>
</w:document>"""


# ---------------------------------------------------------------------------
# Main exporter class
=======
STATUS_COLORS: dict[str, RGBColor] = {
    "pending": RGBColor(0xFF, 0xC1, 0x07),
    "in_progress": RGBColor(0x0D, 0xCA, 0xF0),
    "completed": RGBColor(0x19, 0x87, 0x54),
    "blocked": RGBColor(0xDC, 0x35, 0x45),
    "skipped": RGBColor(0x6C, 0x75, 0x7D),
}

PRIORITY_COLORS: dict[str, RGBColor] = {
    "critical": RGBColor(0xDC, 0x35, 0x45),
    "high": RGBColor(0xFF, 0x85, 0x00),
    "medium": RGBColor(0xFF, 0xC1, 0x07),
    "low": RGBColor(0x19, 0x87, 0x54),
}

DEFAULT_STYLE_CONFIG: dict[str, Any] = {
    "font_name": "Calibri",
    "heading_color": RGBColor(0x0D, 0x6E, 0xFD),
    "header_bg": RGBColor(0x34, 0x3A, 0x40),
    "header_text": RGBColor(0xFF, 0xFF, 0xFF),
    "stripe_bg": RGBColor(0xF2, 0xF2, 0xF2),
    "border_color": RGBColor(0xDE, 0xE2, 0xE6),
}


# ---------------------------------------------------------------------------
# Helper: table cell shading
# ---------------------------------------------------------------------------


def _set_cell_shading(cell: Any, color: RGBColor) -> None:
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), f"{color}")
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_text(cell: Any, text: str, bold: bool = False,
                   color: RGBColor | None = None,
                   font_size: int | None = None) -> None:
    """Set cell text with optional formatting."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if font_size:
        run.font.size = Pt(font_size)


def _format_header_row(row: Any, style_config: dict[str, Any] | None = None) -> None:
    """Format a table header row with background and white text."""
    cfg = style_config or DEFAULT_STYLE_CONFIG
    header_bg = cfg.get("header_bg", DEFAULT_STYLE_CONFIG["header_bg"])
    header_text = cfg.get("header_text", DEFAULT_STYLE_CONFIG["header_text"])

    for cell in row.cells:
        _set_cell_shading(cell, header_bg)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = header_text
                run.bold = True


def _add_alternating_rows(table: Any, style_config: dict[str, Any] | None = None) -> None:
    """Apply alternating row colors to a table (skip header row)."""
    cfg = style_config or DEFAULT_STYLE_CONFIG
    stripe = cfg.get("stripe_bg", DEFAULT_STYLE_CONFIG["stripe_bg"])

    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        if i % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, stripe)


# ---------------------------------------------------------------------------
# Main exporter
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
# ---------------------------------------------------------------------------


class DOCXExporter(TargetExporter):
<<<<<<< HEAD
    """Export execution plans as editable Microsoft Word DOCX documents."""
=======
    """Export execution plans as editable Word DOCX documents."""
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM

    def get_format(self) -> str:
        return "docx"

    def get_extension(self) -> str:
        return ".docx"

    def export(
        self,
        execution_plan: dict[str, Any],
        implementation_brief: dict[str, Any],
        output_path: str,
    ) -> str:
        execution_plan, implementation_brief = self.validate_export_payload(
<<<<<<< HEAD
            execution_plan, implementation_brief
        )
        self.ensure_output_dir(output_path)

        docx_bytes = self.export_plan(execution_plan, brief=implementation_brief)
        with open(output_path, "wb") as f:
            f.write(docx_bytes)
=======
            execution_plan,
            implementation_brief,
        )
        self.ensure_output_dir(output_path)

        doc_bytes = self.export_plan(execution_plan, brief=implementation_brief)
        with open(output_path, "wb") as f:
            f.write(doc_bytes)
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
        return output_path

    def export_plan(
        self,
        plan: dict[str, Any],
<<<<<<< HEAD
        brief: dict[str, Any] | None = None,
        style: str = "default",
    ) -> bytes:
        """Generate a DOCX document for the plan.

        Args:
            plan: Execution plan dictionary.
            brief: Optional implementation brief dictionary.
            style: Style configuration name (default, corporate, modern).
=======
        template: str | None = None,
        brief: dict[str, Any] | None = None,
        style_config: dict[str, Any] | None = None,
    ) -> bytes:
        """Generate a Word DOCX document for the plan.

        Args:
            plan: Execution plan dictionary.
            template: Unused string template name (kept for API symmetry).
            brief: Optional implementation brief dictionary.
            style_config: Optional style overrides.
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM

        Returns:
            DOCX file content as bytes.
        """
<<<<<<< HEAD
        brief = brief or {}
        style_config = STYLE_CONFIGS.get(style, STYLE_CONFIGS["default"])

        body_content = _build_document_body(plan, brief, style_config)
        doc_xml = _document_xml(body_content)

        plan_id = plan.get("id") or "Execution Plan"
        brief_title = brief.get("title") or "Blueprint"

        return _build_docx_zip(
            doc_xml,
            style_config,
            title=plan_id,
            author=brief_title,
            subject="Execution Plan Export",
            keywords="plan, tasks, execution",
        )
=======
        doc = Document()
        brief = brief or {}
        cfg = style_config or DEFAULT_STYLE_CONFIG

        self.apply_styles(doc, cfg)
        self._add_document_properties(doc, plan, brief)
        self._add_cover_page(doc, plan, brief)
        self._add_toc_field(doc)
        self._add_executive_summary(doc, plan, brief)

        tasks = plan.get("tasks") or []
        self.add_task_table(doc, tasks, cfg)
        self.add_timeline_chart(doc, plan)
        self._add_dependency_matrix(doc, tasks)
        self._add_risk_register(doc, tasks)
        self._add_resource_allocation(doc, tasks)
        self._add_appendix(doc, tasks)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM

    def export_with_template(
        self,
        plan: dict[str, Any],
<<<<<<< HEAD
        template_bytes: bytes,
        brief: dict[str, Any] | None = None,
    ) -> bytes:
        """Export using a .dotx template as base.

        Note: Full template processing requires python-docx. This implementation
        generates a standalone DOCX and ignores the template bytes.

        Args:
            plan: Execution plan dictionary.
            template_bytes: Template file bytes (currently unused).
            brief: Optional brief dictionary.
=======
        docx_template: str | bytes,
    ) -> bytes:
        """Generate DOCX using an existing .dotx / .docx template.

        Loads the template, maps plan data into content controls where
        available, and falls back to appending sections.

        Args:
            plan: Execution plan dictionary.
            docx_template: Path to a .dotx/.docx template file, or raw bytes.
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM

        Returns:
            DOCX file content as bytes.
        """
<<<<<<< HEAD
        # Template support requires python-docx for full feature parity;
        # we generate standalone DOCX here.
        _ = template_bytes
        return self.export_plan(plan, brief=brief)

    def add_task_table(
        self,
        tasks: list[dict[str, Any]],
        style: str = "default",
    ) -> str:
        """Generate task table XML fragment.

        Args:
            tasks: List of task dictionaries.
            style: Style name.

        Returns:
            OOXML table fragment string.
        """
        style_config = STYLE_CONFIGS.get(style, STYLE_CONFIGS["default"])
        return _build_task_table(tasks, style_config)

    def apply_styles(self, style_name: str) -> dict[str, str]:
        """Get style configuration by name.

        Args:
            style_name: Style name (default, corporate, modern).

        Returns:
            Style configuration dictionary.
        """
        return STYLE_CONFIGS.get(style_name, STYLE_CONFIGS["default"])


def _build_docx_zip(
    document_xml: str,
    style_config: dict[str, str],
    title: str = "",
    author: str = "",
    subject: str = "",
    keywords: str = "",
) -> bytes:
    """Package all OOXML parts into a ZIP-based .docx file."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", _content_types_xml())
        zf.writestr("_rels/.rels", _rels_xml())
        zf.writestr("word/_rels/document.xml.rels", _word_rels_xml())
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/styles.xml", _styles_xml(style_config))
        zf.writestr("word/numbering.xml", _numbering_xml())
        zf.writestr(
            "docProps/core.xml",
            _core_props_xml(
                title=title, author=author, subject=subject, keywords=keywords
            ),
        )
    return buf.getvalue()
=======
        if isinstance(docx_template, (str, Path)):
            doc = Document(str(docx_template))
        else:
            doc = Document(io.BytesIO(docx_template))

        tasks = plan.get("tasks") or []

        self._map_content_controls(doc, plan)
        self.add_task_table(doc, tasks)
        self.add_timeline_chart(doc, plan)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Public helpers
    # ------------------------------------------------------------------

    def add_task_table(
        self,
        doc: Document,
        tasks: list[dict[str, Any]],
        style_config: dict[str, Any] | None = None,
    ) -> None:
        """Add the task breakdown table to the document."""
        doc.add_page_break()
        doc.add_heading("Task Breakdown", level=1)

        if not tasks:
            doc.add_paragraph("No tasks defined.")
            return

        headers = ["ID", "Title", "Status", "Assignee", "Due Date", "Priority"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            _set_cell_text(cell, header, bold=True)

        _format_header_row(table.rows[0], style_config)

        # Data rows
        for task in tasks:
            row = table.add_row()
            task_id = task.get("id") or ""
            title = task.get("title") or ""
            status = task.get("status") or "pending"
            assignee = task.get("owner_type") or ""
            due_date = task.get("due_date") or ""
            priority = task.get("risk_level") or ""

            row.cells[0].text = task_id
            row.cells[1].text = title

            status_label = STATUS_LABELS.get(status, status)
            _set_cell_text(row.cells[2], status_label)
            status_color = STATUS_COLORS.get(status)
            if status_color:
                _set_cell_shading(row.cells[2], status_color)
                if status in ("completed", "blocked"):
                    for run in row.cells[2].paragraphs[0].runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

            row.cells[3].text = assignee
            row.cells[4].text = str(due_date) if due_date else ""

            if priority:
                _set_cell_text(row.cells[5], priority.capitalize())
                p_color = PRIORITY_COLORS.get(priority.lower())
                if p_color:
                    for run in row.cells[5].paragraphs[0].runs:
                        run.font.color.rgb = p_color

        _add_alternating_rows(table, style_config)

    def add_timeline_chart(
        self,
        doc: Document,
        plan: dict[str, Any],
    ) -> None:
        """Add a timeline visualization section.

        Since python-docx cannot natively embed Excel chart objects,
        this renders a text-based timeline as a table. For true editable
        charts, use export_with_template with a pre-built chart template.
        """
        doc.add_page_break()
        doc.add_heading("Timeline", level=1)

        tasks = plan.get("tasks") or []
        milestones = plan.get("milestones") or []

        if not tasks and not milestones:
            doc.add_paragraph("No timeline data available.")
            return

        # Milestone summary
        if milestones:
            doc.add_heading("Milestones", level=2)
            for ms in milestones:
                name = ms.get("name") or "Milestone"
                desc = ms.get("description") or ""
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(name)
                run.bold = True
                if desc:
                    p.add_run(f" — {desc}")

        # Task timeline table
        if tasks:
            doc.add_heading("Task Schedule", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, h in enumerate(["ID", "Title", "Status", "Milestone"]):
                _set_cell_text(table.rows[0].cells[i], h, bold=True)
            _format_header_row(table.rows[0])

            for task in tasks:
                row = table.add_row()
                row.cells[0].text = task.get("id") or ""
                row.cells[1].text = task.get("title") or ""
                status = task.get("status") or "pending"
                row.cells[2].text = STATUS_LABELS.get(status, status)
                row.cells[3].text = task.get("milestone") or ""

            _add_alternating_rows(table)

    def apply_styles(
        self,
        doc: Document,
        style_config: dict[str, Any] | None = None,
    ) -> None:
        """Apply document-wide styles (font, heading colors)."""
        cfg = style_config or DEFAULT_STYLE_CONFIG
        font_name = cfg.get("font_name", "Calibri")
        heading_color = cfg.get("heading_color", DEFAULT_STYLE_CONFIG["heading_color"])

        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(11)

        for level in range(1, 4):
            heading_style_name = f"Heading {level}"
            if heading_style_name in doc.styles:
                h_style = doc.styles[heading_style_name]
                h_style.font.name = font_name
                h_style.font.color.rgb = heading_color

    # ------------------------------------------------------------------
    # Private section builders
    # ------------------------------------------------------------------

    def _add_document_properties(
        self, doc: Document, plan: dict[str, Any], brief: dict[str, Any]
    ) -> None:
        """Set document core properties."""
        props = doc.core_properties
        plan_id = plan.get("id") or "Execution Plan"
        props.title = f"Execution Plan: {plan_id}"
        props.subject = brief.get("problem_statement") or ""
        props.author = "Blueprint"
        props.keywords = ",".join([
            plan.get("target_engine") or "",
            plan.get("project_type") or "",
        ])
        props.category = "Execution Plan"

    def _add_cover_page(
        self, doc: Document, plan: dict[str, Any], brief: dict[str, Any]
    ) -> None:
        """Add a cover page with title, author, and date."""
        # Spacer
        for _ in range(6):
            doc.add_paragraph("")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"Execution Plan: {plan.get('id') or ''}")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)

        if brief.get("title"):
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = subtitle.add_run(brief["title"])
            sub_run.font.size = Pt(16)
            sub_run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parts = []
        if plan.get("target_engine"):
            parts.append(f"Engine: {plan['target_engine']}")
        if plan.get("target_repo"):
            parts.append(f"Repo: {plan['target_repo']}")
        parts.append(f"Generated: {datetime.now().strftime('%Y-%m-%d')}")
        meta.add_run(" | ".join(parts)).font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

        doc.add_paragraph("")
        doc.add_paragraph("")

        author_p = doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_p.add_run("Author: Blueprint").font.size = Pt(12)

    def _add_toc_field(self, doc: Document) -> None:
        """Insert an auto-updating Table of Contents field."""
        doc.add_page_break()
        doc.add_heading("Table of Contents", level=1)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar)

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run._r.append(instrText)

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")
        run._r.append(fldChar2)

        # Placeholder text
        run2 = paragraph.add_run("[Update this field to generate Table of Contents]")
        run2.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)
        run2.font.italic = True

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run3 = paragraph.add_run()
        run3._r.append(fldChar3)

    def _add_executive_summary(
        self, doc: Document, plan: dict[str, Any], brief: dict[str, Any]
    ) -> None:
        """Add the executive summary section."""
        doc.add_page_break()
        doc.add_heading("Executive Summary", level=1)

        tasks = plan.get("tasks") or []
        total = len(tasks)
        completed = sum(1 for t in tasks if t.get("status") == "completed")
        progress = round(completed / total * 100) if total > 0 else 0

        status = plan.get("status") or "draft"
        doc.add_paragraph(
            f"Plan Status: {STATUS_LABELS.get(status, status)}"
        )
        doc.add_paragraph(
            f"Progress: {progress}% ({completed}/{total} tasks completed)"
        )

        if brief.get("problem_statement"):
            doc.add_heading("Problem Statement", level=2)
            doc.add_paragraph(str(brief["problem_statement"]))

        if brief.get("mvp_goal"):
            doc.add_heading("MVP Goal", level=2)
            doc.add_paragraph(str(brief["mvp_goal"]))

        if brief.get("risks"):
            doc.add_heading("Key Risks", level=2)
            for risk in brief["risks"]:
                doc.add_paragraph(str(risk), style="List Bullet")

        # Status summary table
        if tasks:
            doc.add_heading("Status Overview", level=2)
            status_counts = Counter(t.get("status") or "pending" for t in tasks)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            _set_cell_text(table.rows[0].cells[0], "Status", bold=True)
            _set_cell_text(table.rows[0].cells[1], "Count", bold=True)
            _format_header_row(table.rows[0])

            for s_key in ["completed", "in_progress", "pending", "blocked", "skipped"]:
                count = status_counts.get(s_key, 0)
                if count > 0:
                    row = table.add_row()
                    _set_cell_text(row.cells[0], STATUS_LABELS.get(s_key, s_key))
                    row.cells[1].text = str(count)
                    s_color = STATUS_COLORS.get(s_key)
                    if s_color:
                        _set_cell_shading(row.cells[0], s_color)
                        if s_key in ("completed", "blocked"):
                            for run in row.cells[0].paragraphs[0].runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def _add_dependency_matrix(
        self, doc: Document, tasks: list[dict[str, Any]]
    ) -> None:
        """Add a dependency matrix table."""
        doc.add_page_break()
        doc.add_heading("Dependency Matrix", level=1)

        deps_exist = any(task.get("depends_on") for task in tasks)
        if not deps_exist:
            doc.add_paragraph("No dependencies defined.")
            return

        task_map = {t.get("id") or "": t for t in tasks}
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        for i, h in enumerate(["Task", "Depends On", "Status"]):
            _set_cell_text(table.rows[0].cells[i], h, bold=True)
        _format_header_row(table.rows[0])

        for task in tasks:
            deps = task.get("depends_on") or []
            if not deps:
                continue
            for dep_id in deps:
                row = table.add_row()
                task_id = task.get("id") or ""
                task_title = task.get("title") or ""
                row.cells[0].text = f"{task_id}: {task_title}"

                dep_task = task_map.get(dep_id, {})
                dep_title = dep_task.get("title") or str(dep_id)
                row.cells[1].text = f"{dep_id}: {dep_title}"

                dep_status = dep_task.get("status") or "unknown"
                row.cells[2].text = STATUS_LABELS.get(dep_status, dep_status)

        _add_alternating_rows(table)

    def _add_risk_register(
        self, doc: Document, tasks: list[dict[str, Any]]
    ) -> None:
        """Add a risk register table."""
        doc.add_page_break()
        doc.add_heading("Risk Register", level=1)

        risk_tasks = [t for t in tasks if t.get("risk_level")]
        if not risk_tasks:
            doc.add_paragraph("No risk levels assigned.")
            return

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        for i, h in enumerate(["ID", "Title", "Risk Level", "Complexity"]):
            _set_cell_text(table.rows[0].cells[i], h, bold=True)
        _format_header_row(table.rows[0])

        for task in risk_tasks:
            row = table.add_row()
            row.cells[0].text = task.get("id") or ""
            row.cells[1].text = task.get("title") or ""

            risk = (task.get("risk_level") or "").lower()
            _set_cell_text(row.cells[2], risk.capitalize())
            r_color = PRIORITY_COLORS.get(risk)
            if r_color:
                for run in row.cells[2].paragraphs[0].runs:
                    run.font.color.rgb = r_color

            row.cells[3].text = task.get("estimated_complexity") or ""

        _add_alternating_rows(table)

    def _add_resource_allocation(
        self, doc: Document, tasks: list[dict[str, Any]]
    ) -> None:
        """Add resource allocation table grouped by owner type."""
        doc.add_page_break()
        doc.add_heading("Resource Allocation", level=1)

        owner_tasks: dict[str, list[dict[str, Any]]] = {}
        for task in tasks:
            owner = task.get("owner_type") or "unassigned"
            owner_tasks.setdefault(owner, []).append(task)

        if not owner_tasks:
            doc.add_paragraph("No resource data available.")
            return

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        for i, h in enumerate(["Owner", "Task Count", "Est. Hours", "Completed"]):
            _set_cell_text(table.rows[0].cells[i], h, bold=True)
        _format_header_row(table.rows[0])

        for owner, o_tasks in sorted(owner_tasks.items()):
            row = table.add_row()
            row.cells[0].text = owner
            row.cells[1].text = str(len(o_tasks))
            total_hours = sum(t.get("estimated_hours") or 0 for t in o_tasks)
            row.cells[2].text = f"{total_hours:.1f}" if total_hours else "—"
            done = sum(1 for t in o_tasks if t.get("status") == "completed")
            row.cells[3].text = f"{done}/{len(o_tasks)}"

        _add_alternating_rows(table)

    def _add_appendix(
        self, doc: Document, tasks: list[dict[str, Any]]
    ) -> None:
        """Add appendix with detailed task descriptions."""
        doc.add_page_break()
        doc.add_heading("Appendix: Task Details", level=1)

        if not tasks:
            doc.add_paragraph("No tasks to detail.")
            return

        for task in tasks:
            task_id = task.get("id") or ""
            title = task.get("title") or ""
            doc.add_heading(f"{task_id}: {title}", level=2)

            desc = task.get("description") or "No description provided."
            doc.add_paragraph(desc)

            # Acceptance criteria
            criteria = task.get("acceptance_criteria") or []
            if criteria:
                doc.add_heading("Acceptance Criteria", level=3)
                for criterion in criteria:
                    doc.add_paragraph(str(criterion), style="List Bullet")

            # Files
            files = task.get("files_or_modules") or []
            if files:
                doc.add_heading("Files / Modules", level=3)
                for f in files:
                    doc.add_paragraph(str(f), style="List Bullet")

            # Test command
            test_cmd = task.get("test_command")
            if test_cmd:
                doc.add_paragraph(f"Test command: {test_cmd}")

    def _map_content_controls(
        self, doc: Document, plan: dict[str, Any]
    ) -> None:
        """Map plan data to Word content controls (structured document tags).

        Searches for SDT elements with specific aliases and replaces
        their content with plan data. Falls back gracefully if no
        content controls are found.
        """
        mapping = {
            "plan_id": plan.get("id") or "",
            "plan_status": STATUS_LABELS.get(plan.get("status") or "", plan.get("status") or ""),
            "target_engine": plan.get("target_engine") or "",
            "target_repo": plan.get("target_repo") or "",
        }

        body = doc.element.body
        for sdt in body.iter(qn("w:sdt")):
            sdt_pr = sdt.find(qn("w:sdtPr"))
            if sdt_pr is None:
                continue
            alias = sdt_pr.find(qn("w:alias"))
            if alias is None:
                continue
            alias_val = alias.get(qn("w:val"), "")
            if alias_val in mapping:
                sdt_content = sdt.find(qn("w:sdtContent"))
                if sdt_content is not None:
                    for p in sdt_content.findall(qn("w:p")):
                        for r in p.findall(qn("w:r")):
                            for t in r.findall(qn("w:t")):
                                t.text = mapping[alias_val]
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
