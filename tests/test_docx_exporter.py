<<<<<<< HEAD
"""Tests for the Word DOCX exporter with editable tables."""

import io
import os
import zipfile

from blueprint.exporters.docx_exporter import (
    DOCXExporter,
    STYLE_CONFIGS,
    _count_statuses,
=======
"""Tests for the Word DOCX exporter."""

import io
import os

from docx import Document

from blueprint.exporters.docx_exporter import (
    DOCXExporter,
    STATUS_COLORS,
    STATUS_LABELS,
    DEFAULT_STYLE_CONFIG,
    _set_cell_shading,
    _format_header_row,
    _add_alternating_rows,
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _sample_plan() -> dict:
    return {
        "id": "plan-001",
        "implementation_brief_id": "ib-001",
        "target_engine": "relay",
        "target_repo": "acme/widgets",
        "project_type": "web",
        "milestones": [
            {"name": "Sprint 1", "description": "First sprint"},
            {"name": "Sprint 2", "description": "Second sprint"},
        ],
        "test_strategy": "pytest",
        "handoff_prompt": None,
        "status": "in_progress",
        "created_at": None,
        "updated_at": None,
        "generation_model": None,
        "generation_tokens": None,
        "generation_prompt": None,
        "metadata": {},
        "tasks": [
            {
                "id": "task-1",
                "execution_plan_id": "plan-001",
                "title": "Implement auth",
                "description": "Build OAuth2 login flow",
                "milestone": "Sprint 1",
                "owner_type": "developer",
                "suggested_engine": None,
                "depends_on": [],
                "files_or_modules": ["src/auth.py"],
                "acceptance_criteria": ["Users can log in", "OAuth tokens refresh"],
                "estimated_complexity": "high",
                "estimated_hours": 16.0,
                "risk_level": "medium",
                "test_command": "pytest tests/test_auth.py",
                "status": "completed",
                "metadata": {},
                "blocked_reason": None,
                "created_at": None,
                "updated_at": None,
            },
            {
                "id": "task-2",
                "execution_plan_id": "plan-001",
                "title": "Add dashboard",
                "description": "Create user dashboard page",
                "milestone": "Sprint 1",
                "owner_type": "developer",
                "suggested_engine": None,
                "depends_on": ["task-1"],
                "files_or_modules": ["src/dashboard.py"],
                "acceptance_criteria": ["Dashboard shows metrics"],
                "estimated_complexity": "medium",
                "estimated_hours": 8.0,
                "risk_level": "low",
                "test_command": None,
                "status": "in_progress",
                "metadata": {},
                "blocked_reason": None,
                "created_at": None,
                "updated_at": None,
            },
            {
                "id": "task-3",
                "execution_plan_id": "plan-001",
                "title": "Deploy to staging",
                "description": "Deploy application to staging environment",
                "milestone": "Sprint 2",
                "owner_type": "devops",
                "suggested_engine": None,
                "depends_on": ["task-2"],
                "files_or_modules": None,
                "acceptance_criteria": ["App is live on staging"],
                "estimated_complexity": "low",
                "estimated_hours": 2.0,
                "risk_level": "low",
                "test_command": None,
                "status": "pending",
                "metadata": {},
                "blocked_reason": None,
                "created_at": None,
                "updated_at": None,
            },
        ],
    }


def _sample_brief() -> dict:
    return {
        "id": "ib-001",
        "source_brief_id": "sb-001",
        "title": "Widget Authentication System",
        "domain": "web",
        "target_user": "developers",
        "buyer": None,
        "workflow_context": None,
        "problem_statement": "Users need secure authentication for the widget platform",
        "mvp_goal": "Working OAuth2 login",
        "product_surface": "web",
        "scope": ["auth", "dashboard"],
        "non_goals": ["mobile"],
        "assumptions": ["OAuth provider is configured"],
        "architecture_notes": None,
        "data_requirements": None,
        "integration_points": [],
        "risks": ["OAuth provider downtime"],
        "validation_plan": "Integration tests",
        "definition_of_done": ["All tests pass"],
        "status": "planned",
        "created_at": None,
        "updated_at": None,
        "generation_model": None,
        "generation_tokens": None,
        "generation_prompt": None,
    }


<<<<<<< HEAD
# ---------------------------------------------------------------------------
# DOCX generation tests
# ---------------------------------------------------------------------------


def test_export_plan_generates_valid_docx():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    assert isinstance(docx, bytes)
    # DOCX files are ZIP archives
    assert zipfile.is_zipfile(io.BytesIO(docx))


def test_export_plan_contains_required_parts():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        names = zf.namelist()
        assert "[Content_Types].xml" in names
        assert "_rels/.rels" in names
        assert "word/document.xml" in names
        assert "word/styles.xml" in names
        assert "word/numbering.xml" in names
        assert "word/_rels/document.xml.rels" in names
        assert "docProps/core.xml" in names


def test_export_plan_document_xml_valid():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert '<?xml version="1.0"' in doc
        assert "<w:document" in doc
        assert "<w:body>" in doc
        assert "</w:document>" in doc


# ---------------------------------------------------------------------------
# Content tests
# ---------------------------------------------------------------------------


def test_export_plan_contains_cover_page():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Execution Plan" in doc
        assert "plan-001" in doc


def test_export_plan_contains_executive_summary():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Executive Summary" in doc
        assert "Progress:" in doc


def test_export_plan_contains_task_data():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Implement auth" in doc
        assert "Add dashboard" in doc
        assert "Deploy to staging" in doc
        assert "task-1" in doc
        assert "task-2" in doc
        assert "task-3" in doc


def test_export_plan_contains_task_table():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "<w:tbl>" in doc
        assert "Task Breakdown" in doc
        # Table header columns
        assert "Status" in doc
        assert "Milestone" in doc
        assert "Est. Hours" in doc


def test_export_plan_contains_acceptance_criteria():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Users can log in" in doc
        assert "OAuth tokens refresh" in doc


def test_export_plan_contains_dependency_matrix():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Dependency Matrix" in doc


def test_export_plan_contains_risk_register():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Risk Register" in doc
        assert "OAuth provider downtime" in doc


def test_export_plan_contains_resource_allocation():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Resource Allocation" in doc
        assert "developer" in doc
        assert "devops" in doc


def test_export_plan_contains_timeline():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Timeline" in doc
        assert "Sprint 1" in doc
        assert "Sprint 2" in doc


def test_export_plan_contains_appendix():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "Appendix: Detailed Task Descriptions" in doc
        assert "Build OAuth2 login flow" in doc


# ---------------------------------------------------------------------------
# Word styles tests
# ---------------------------------------------------------------------------


def test_export_plan_uses_heading_styles():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert 'w:val="Heading1"' in doc
        assert 'w:val="Heading2"' in doc


def test_styles_xml_contains_heading_definitions():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
        assert 'w:styleId="Heading1"' in styles
        assert 'w:styleId="Heading2"' in styles
        assert 'w:styleId="Heading3"' in styles
        assert 'w:styleId="Normal"' in styles


def test_styles_xml_contains_table_grid():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
        assert 'w:styleId="TableGrid"' in styles


def test_table_has_formatting():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        # Alternating row colors (shading)
        assert "w:shd" in doc
        # Table borders
        assert "w:tblBorders" in doc


def test_status_colors_in_table():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        # Status colors should be present
        assert "008000" in doc  # completed green
        assert "0066CC" in doc or "CC9900" in doc  # in_progress blue or pending yellow


# ---------------------------------------------------------------------------
# Style configuration tests
# ---------------------------------------------------------------------------


def test_default_style():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), style="default")
    assert isinstance(docx, bytes)
    assert zipfile.is_zipfile(io.BytesIO(docx))


def test_corporate_style():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), style="corporate")
    assert isinstance(docx, bytes)

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
        assert "1B365D" in styles  # corporate primary color


def test_modern_style():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), style="modern")
    assert isinstance(docx, bytes)

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        styles = zf.read("word/styles.xml").decode("utf-8")
        assert "2D3748" in styles  # modern primary color


def test_unknown_style_falls_back():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), style="nonexistent")
    assert isinstance(docx, bytes)
    assert zipfile.is_zipfile(io.BytesIO(docx))


def test_style_configs_have_required_keys():
    required_keys = {"primary", "header_bg", "header_text", "stripe_bg", "border"}
    for name, config in STYLE_CONFIGS.items():
        missing = required_keys - set(config.keys())
        assert not missing, f"Style '{name}' missing keys: {missing}"


def test_apply_styles():
    exporter = DOCXExporter()
    config = exporter.apply_styles("corporate")
    assert config["primary"] == "1B365D"


def test_apply_styles_unknown():
    exporter = DOCXExporter()
    config = exporter.apply_styles("unknown")
    assert config == STYLE_CONFIGS["default"]


# ---------------------------------------------------------------------------
# Editable task table tests
# ---------------------------------------------------------------------------


def test_add_task_table():
    exporter = DOCXExporter()
    tasks = _sample_plan()["tasks"]
    xml = exporter.add_task_table(tasks)

    assert "<w:tbl>" in xml
    assert "Implement auth" in xml
    assert "Add dashboard" in xml
    assert "Deploy to staging" in xml


def test_task_table_has_sortable_columns():
    exporter = DOCXExporter()
    tasks = _sample_plan()["tasks"]
    xml = exporter.add_task_table(tasks)

    # All column headers present
    for header in ["ID", "Title", "Status", "Milestone", "Owner", "Complexity", "Est. Hours"]:
        assert header in xml


# ---------------------------------------------------------------------------
# Template support
# ---------------------------------------------------------------------------


def test_export_with_template():
    exporter = DOCXExporter()
    template_bytes = b"fake template data"
    docx = exporter.export_with_template(_sample_plan(), template_bytes, brief=_sample_brief())

    assert isinstance(docx, bytes)
    assert zipfile.is_zipfile(io.BytesIO(docx))


# ---------------------------------------------------------------------------
# Metadata tests
# ---------------------------------------------------------------------------


def test_core_properties():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan(), brief=_sample_brief())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        core = zf.read("docProps/core.xml").decode("utf-8")
        assert "plan-001" in core  # title
        assert "Widget Authentication System" in core  # author
        assert "dcterms:created" in core
        assert "dcterms:modified" in core


# ---------------------------------------------------------------------------
# File export test
# ---------------------------------------------------------------------------


def test_export_to_file(tmp_path):
    exporter = DOCXExporter()
    output_path = str(tmp_path / "plan.docx")

    result = exporter.export(_sample_plan(), _sample_brief(), output_path)

    assert result == output_path
    assert os.path.exists(output_path)

    with open(output_path, "rb") as f:
        content = f.read()

    assert zipfile.is_zipfile(io.BytesIO(content))


def test_export_creates_directory(tmp_path):
    exporter = DOCXExporter()
    output_path = str(tmp_path / "subdir" / "deep" / "plan.docx")

    result = exporter.export(_sample_plan(), _sample_brief(), output_path)

    assert result == output_path
    assert os.path.exists(output_path)


# ---------------------------------------------------------------------------
# Format and extension tests
# ---------------------------------------------------------------------------


def test_export_format_and_extension():
=======
def _load_docx(data: bytes) -> Document:
    """Load a Document from raw bytes."""
    return Document(io.BytesIO(data))


def _all_text(doc: Document) -> str:
    """Extract all text from a Document."""
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Tests: basic generation
# ---------------------------------------------------------------------------


def test_export_plan_returns_bytes():
    exporter = DOCXExporter()
    result = exporter.export_plan(_sample_plan())
    assert isinstance(result, bytes)
    assert len(result) > 0


def test_export_plan_is_valid_docx():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    doc = _load_docx(data)
    assert len(doc.paragraphs) > 0


def test_format_and_extension():
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
    exporter = DOCXExporter()
    assert exporter.get_format() == "docx"
    assert exporter.get_extension() == ".docx"


# ---------------------------------------------------------------------------
<<<<<<< HEAD
# Empty plan handling
=======
# Tests: cover page
# ---------------------------------------------------------------------------


def test_cover_page_contains_plan_id():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "plan-001" in text


def test_cover_page_contains_brief_title():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan(), brief=_sample_brief())
    text = _all_text(_load_docx(data))
    assert "Widget Authentication System" in text


# ---------------------------------------------------------------------------
# Tests: executive summary
# ---------------------------------------------------------------------------


def test_executive_summary_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan(), brief=_sample_brief())
    text = _all_text(_load_docx(data))
    assert "Executive Summary" in text


def test_executive_summary_progress():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    # 1 completed out of 3 = 33%
    assert "33%" in text


def test_executive_summary_problem_statement():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan(), brief=_sample_brief())
    text = _all_text(_load_docx(data))
    assert "secure authentication" in text


def test_executive_summary_risks():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan(), brief=_sample_brief())
    text = _all_text(_load_docx(data))
    assert "OAuth provider downtime" in text


def test_executive_summary_status_overview():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Completed" in text
    assert "In Progress" in text
    assert "Pending" in text


# ---------------------------------------------------------------------------
# Tests: task table
# ---------------------------------------------------------------------------


def test_task_table_contains_all_tasks():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Implement auth" in text
    assert "Add dashboard" in text
    assert "Deploy to staging" in text


def test_task_table_has_headers():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Task Breakdown" in text
    for header in ["ID", "Title", "Status", "Assignee"]:
        assert header in text


def test_task_table_uses_table_grid_style():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    doc = _load_docx(data)
    table_styles = [t.style.name for t in doc.tables if t.style]
    assert "Table Grid" in table_styles


def test_task_table_status_labels():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Completed" in text
    assert "In Progress" in text
    assert "Pending" in text


# ---------------------------------------------------------------------------
# Tests: timeline
# ---------------------------------------------------------------------------


def test_timeline_section_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Timeline" in text


def test_timeline_contains_milestones():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Sprint 1" in text
    assert "Sprint 2" in text


def test_timeline_contains_task_schedule():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Task Schedule" in text


# ---------------------------------------------------------------------------
# Tests: dependency matrix
# ---------------------------------------------------------------------------


def test_dependency_matrix_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Dependency Matrix" in text


def test_dependency_matrix_shows_deps():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    # task-2 depends on task-1
    assert "task-1" in text
    assert "task-2" in text


# ---------------------------------------------------------------------------
# Tests: risk register
# ---------------------------------------------------------------------------


def test_risk_register_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Risk Register" in text


def test_risk_register_shows_risk_levels():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Medium" in text
    assert "Low" in text


# ---------------------------------------------------------------------------
# Tests: resource allocation
# ---------------------------------------------------------------------------


def test_resource_allocation_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Resource Allocation" in text


def test_resource_allocation_owners():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "developer" in text
    assert "devops" in text


# ---------------------------------------------------------------------------
# Tests: appendix
# ---------------------------------------------------------------------------


def test_appendix_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Appendix" in text


def test_appendix_descriptions():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Build OAuth2 login flow" in text
    assert "Create user dashboard page" in text


def test_appendix_acceptance_criteria():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Users can log in" in text
    assert "OAuth tokens refresh" in text


def test_appendix_files():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "src/auth.py" in text


def test_appendix_test_command():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "pytest tests/test_auth.py" in text


# ---------------------------------------------------------------------------
# Tests: Word styles
# ---------------------------------------------------------------------------


def test_heading_styles_applied():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    doc = _load_docx(data)
    heading_styles = set()
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            heading_styles.add(p.style.name)
    assert "Heading 1" in heading_styles
    assert "Heading 2" in heading_styles


def test_normal_style_font():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    doc = _load_docx(data)
    normal = doc.styles["Normal"]
    assert normal.font.name == "Calibri"


def test_custom_style_config():
    from docx.shared import RGBColor as RC

    exporter = DOCXExporter()
    custom_cfg = {
        "font_name": "Arial",
        "heading_color": RC(0xFF, 0x00, 0x00),
        "header_bg": RC(0x00, 0x00, 0x00),
        "header_text": RC(0xFF, 0xFF, 0xFF),
        "stripe_bg": RC(0xEE, 0xEE, 0xEE),
    }
    data = exporter.export_plan(_sample_plan(), style_config=custom_cfg)
    doc = _load_docx(data)
    assert doc.styles["Normal"].font.name == "Arial"


# ---------------------------------------------------------------------------
# Tests: document properties
# ---------------------------------------------------------------------------


def test_document_properties():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan(), brief=_sample_brief())
    doc = _load_docx(data)
    props = doc.core_properties
    assert "plan-001" in props.title
    assert props.author == "Blueprint"
    assert "secure authentication" in props.subject


# ---------------------------------------------------------------------------
# Tests: table of contents
# ---------------------------------------------------------------------------


def test_toc_field_present():
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    text = _all_text(_load_docx(data))
    assert "Table of Contents" in text


# ---------------------------------------------------------------------------
# Tests: file export
# ---------------------------------------------------------------------------


def test_export_to_file(tmp_path):
    exporter = DOCXExporter()
    output_path = str(tmp_path / "plan.docx")
    result = exporter.export(_sample_plan(), _sample_brief(), output_path)
    assert result == output_path
    assert os.path.exists(output_path)

    doc = Document(output_path)
    text = _all_text(doc)
    assert "plan-001" in text


# ---------------------------------------------------------------------------
# Tests: empty plan
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
# ---------------------------------------------------------------------------


def test_empty_plan():
    exporter = DOCXExporter()
    plan = {
        "id": "plan-empty",
        "implementation_brief_id": "ib-001",
        "target_engine": None,
        "target_repo": None,
        "project_type": None,
        "milestones": [],
        "test_strategy": None,
        "handoff_prompt": None,
        "status": "draft",
        "created_at": None,
        "updated_at": None,
        "generation_model": None,
        "generation_tokens": None,
        "generation_prompt": None,
        "metadata": {},
        "tasks": [],
    }
<<<<<<< HEAD

    docx = exporter.export_plan(plan)
    assert isinstance(docx, bytes)
    assert zipfile.is_zipfile(io.BytesIO(docx))

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "plan-empty" in doc
        assert "0%" in doc


# ---------------------------------------------------------------------------
# Utility function tests
# ---------------------------------------------------------------------------


def test_count_statuses():
    tasks = [
        {"status": "completed"},
        {"status": "completed"},
        {"status": "in_progress"},
        {"status": "pending"},
        {"status": "blocked"},
    ]
    counts = _count_statuses(tasks)
    assert counts["completed"] == 2
    assert counts["in_progress"] == 1
    assert counts["pending"] == 1
    assert counts["blocked"] == 1


def test_count_statuses_empty():
    assert _count_statuses([]) == {}


# ---------------------------------------------------------------------------
# Security / XSS prevention
# ---------------------------------------------------------------------------


def test_xml_escaping():
    exporter = DOCXExporter()
    plan = _sample_plan()
    plan["tasks"][0]["title"] = '<script>alert("xss")</script>'

    docx = exporter.export_plan(plan)

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert '<script>alert("xss")</script>' not in doc
        assert "&lt;script&gt;" in doc


# ---------------------------------------------------------------------------
# Dependency table tests
# ---------------------------------------------------------------------------


def test_dependency_table_with_deps():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        # task-2 depends on task-1
        assert "Depends On" in doc or "Source Task" in doc


def test_dependency_table_no_deps():
    plan = _sample_plan()
    for task in plan["tasks"]:
        task["depends_on"] = []

    exporter = DOCXExporter()
    docx = exporter.export_plan(plan)

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert "No dependencies defined" in doc


# ---------------------------------------------------------------------------
# Page layout tests
# ---------------------------------------------------------------------------


def test_page_layout():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        # US Letter page size (12240 x 15840 twips)
        assert "w:w=\"12240\"" in doc
        assert "w:h=\"15840\"" in doc
        # 1-inch margins (1440 twips)
        assert "w:top=\"1440\"" in doc


def test_page_breaks():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        doc = zf.read("word/document.xml").decode("utf-8")
        assert 'w:type="page"' in doc


# ---------------------------------------------------------------------------
# Numbering (list) support
# ---------------------------------------------------------------------------


def test_numbering_xml():
    exporter = DOCXExporter()
    docx = exporter.export_plan(_sample_plan())

    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        numbering = zf.read("word/numbering.xml").decode("utf-8")
        assert "w:abstractNum" in numbering
        assert "w:num" in numbering
=======
    data = exporter.export_plan(plan)
    assert isinstance(data, bytes)
    text = _all_text(_load_docx(data))
    assert "plan-empty" in text
    assert "No tasks defined" in text


# ---------------------------------------------------------------------------
# Tests: template support
# ---------------------------------------------------------------------------


def test_export_with_template_from_bytes():
    exporter = DOCXExporter()
    # Create a minimal template document
    template_doc = Document()
    template_doc.add_heading("Template Header", level=1)
    buf = io.BytesIO()
    template_doc.save(buf)
    template_bytes = buf.getvalue()

    result = exporter.export_with_template(_sample_plan(), template_bytes)
    assert isinstance(result, bytes)
    doc = _load_docx(result)
    text = _all_text(doc)
    assert "Template Header" in text
    assert "Task Breakdown" in text


def test_export_with_template_from_path(tmp_path):
    exporter = DOCXExporter()
    template_path = tmp_path / "template.docx"
    template_doc = Document()
    template_doc.add_heading("My Template", level=1)
    template_doc.save(str(template_path))

    result = exporter.export_with_template(_sample_plan(), str(template_path))
    assert isinstance(result, bytes)
    doc = _load_docx(result)
    text = _all_text(doc)
    assert "My Template" in text


# ---------------------------------------------------------------------------
# Tests: helper functions
# ---------------------------------------------------------------------------


def test_set_cell_shading():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    from docx.shared import RGBColor as RC
    _set_cell_shading(cell, RC(0xFF, 0x00, 0x00))
    # Verify shading element was added
    tc_pr = cell._tc.tcPr
    assert tc_pr is not None


def test_format_header_row():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Col A"
    table.rows[0].cells[1].text = "Col B"
    _format_header_row(table.rows[0])
    # Verify formatting was applied (cells have tcPr with shading)
    for cell in table.rows[0].cells:
        assert cell._tc.tcPr is not None


def test_add_alternating_rows():
    doc = Document()
    table = doc.add_table(rows=4, cols=2)
    for i in range(4):
        table.rows[i].cells[0].text = f"Row {i}"
    _add_alternating_rows(table)
    # Row 0 (header) should not be striped, row 2 (even) should
    assert table.rows[2].cells[0]._tc.tcPr is not None


# ---------------------------------------------------------------------------
# Tests: accessibility
# ---------------------------------------------------------------------------


def test_heading_hierarchy():
    """Verify proper heading hierarchy: H1 before H2 before H3."""
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan(), brief=_sample_brief())
    doc = _load_docx(data)

    heading_levels = []
    for p in doc.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            level = int(p.style.name.split()[-1])
            heading_levels.append(level)

    assert len(heading_levels) > 0
    # First heading should be level 1
    assert heading_levels[0] == 1
    # No heading should skip more than one level
    for i in range(1, len(heading_levels)):
        assert heading_levels[i] <= heading_levels[i - 1] + 1 or heading_levels[i] == 1


def test_table_headers_marked():
    """Verify that tables have header rows with bold text."""
    exporter = DOCXExporter()
    data = exporter.export_plan(_sample_plan())
    doc = _load_docx(data)

    for table in doc.tables:
        if len(table.rows) > 1:
            header_row = table.rows[0]
            has_bold = False
            for cell in header_row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.bold:
                            has_bold = True
                            break
            assert has_bold, "Table header row should have bold text"


# ---------------------------------------------------------------------------
# Tests: constants
# ---------------------------------------------------------------------------


def test_status_labels_complete():
    expected = {"pending", "in_progress", "completed", "blocked", "skipped"}
    assert set(STATUS_LABELS.keys()) == expected


def test_status_colors_complete():
    expected = {"pending", "in_progress", "completed", "blocked", "skipped"}
    assert set(STATUS_COLORS.keys()) == expected


def test_default_style_config_keys():
    required = {"font_name", "heading_color", "header_bg", "header_text", "stripe_bg"}
    assert required.issubset(set(DEFAULT_STYLE_CONFIG.keys()))
>>>>>>> relay/claude-code/add-plan-export-to-word-docx-with-editable-tables-01KR6YNM
